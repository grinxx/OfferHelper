#!/usr/bin/env python3
"""Extract resume text from PDF, DOCX, TXT, or Markdown files."""

from __future__ import annotations

import argparse
import pathlib
import sys


def extract_pdf(path: pathlib.Path) -> str:
    try:
        import pdfplumber
    except Exception as exc:  # pragma: no cover - environment dependent
        raise RuntimeError("pdfplumber is required to read PDF files") from exc

    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
    return "\n\n".join(parts)


def extract_docx(path: pathlib.Path) -> str:
    try:
        import docx
    except Exception as exc:  # pragma: no cover - environment dependent
        raise RuntimeError("python-docx is required to read DOCX files") from exc

    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                table_cells.append(" | ".join(values))
    return "\n".join(paragraphs + table_cells)


def extract_text(path: pathlib.Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in {".txt", ".md", ".markdown"}:
        return path.read_text(encoding="utf-8")
    raise ValueError(f"Unsupported resume file type: {suffix}")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, help="Resume file path")
    parser.add_argument("--output", help="Optional output .txt path")
    args = parser.parse_args()

    input_path = pathlib.Path(args.input)
    if not input_path.exists():
        print(f"Input file not found: {input_path}", file=sys.stderr)
        return 2

    text = extract_text(input_path).strip()
    if args.output:
        output_path = pathlib.Path(args.output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(text + "\n", encoding="utf-8")
    else:
        print(text)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
